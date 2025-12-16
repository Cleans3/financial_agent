"""
Document Service - Handles file uploads, text extraction, and document ingestion
Supports: PDF, DOCX, TXT, PNG, JPG (with OCR)
"""

import logging
import os
import tempfile
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from datetime import datetime

import pdfplumber
import pytesseract
from PIL import Image
import cv2

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from src.services.rag_service import RAGService

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentService:
    """
    Service for document ingestion and processing
    
    Features:
    - Multiple file format support (PDF, DOCX, TXT, Images)
    - OCR for images and scanned PDFs
    - Text extraction and chunking
    - Integration with RAG service
    """
    
    # Supported file types
    SUPPORTED_FORMATS = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg'
    }
    
    # File size limits (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    def __init__(self, rag_service: Optional[RAGService] = None):
        """
        Initialize Document Service
        
        Args:
            rag_service: RAGService instance (optional, creates new if not provided)
        """
        self.rag_service = rag_service or RAGService()
        logger.info("Document Service initialized")
    
    def process_file(self, 
                    file_path: str,
                    doc_id: str,
                    title: str = "",
                    user_id: str = "") -> Tuple[bool, str, int]:
        """
        Process a file and ingest into RAG service
        
        Args:
            file_path: Path to the file
            doc_id: Unique document ID
            title: Document title (auto-generated from filename if not provided)
            user_id: User who uploaded the document
            
        Returns:
            Tuple of (success, message, chunks_added)
        """
        try:
            file_path = Path(file_path)
            
            # Validate file exists
            if not file_path.exists():
                return False, f"File not found: {file_path}", 0
            
            # Validate file size
            file_size = file_path.stat().st_size
            if file_size > self.MAX_FILE_SIZE:
                return False, f"File too large: {file_size / 1024 / 1024:.1f}MB (max 50MB)", 0
            
            # Get file extension
            file_ext = file_path.suffix.lower().lstrip('.')
            if file_ext not in self.SUPPORTED_FORMATS:
                return False, f"Unsupported file format: {file_ext}", 0
            
            # Auto-generate title if not provided
            if not title:
                title = file_path.stem
            
            logger.info(f"Processing file: {file_path.name} ({file_ext})")
            
            # Extract text based on file type
            text = None
            if file_ext == 'pdf':
                text = self._extract_pdf(str(file_path))
            elif file_ext == 'docx':
                text = self._extract_docx(str(file_path))
            elif file_ext == 'txt':
                text = self._extract_txt(str(file_path))
            elif file_ext in ['png', 'jpg', 'jpeg']:
                text = self._extract_image(str(file_path))
            
            if not text or len(text.strip()) == 0:
                return False, "No text extracted from file", 0
            
            logger.info(f"Extracted {len(text)} characters from {file_path.name}")
            
            # Add to RAG service
            chunks_added = self.rag_service.add_document(
                doc_id=doc_id,
                text=text,
                title=title,
                source=file_path.name,
                user_id=user_id
            )
            
            return True, f"Successfully processed {title} ({chunks_added} chunks)", chunks_added
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return False, f"Error processing file: {str(e)}", 0
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF has {total_pages} pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Try text extraction first
                        page_text = page.extract_text()
                        
                        if page_text and len(page_text.strip()) > 0:
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                        else:
                            # If no text, try OCR
                            logger.info(f"Page {page_num} has no text, attempting OCR...")
                            ocr_text = self._ocr_pdf_page(page)
                            if ocr_text:
                                text_parts.append(f"--- Page {page_num} (OCR) ---\n{ocr_text}")
                    except Exception as page_error:
                        logger.warning(f"Error processing page {page_num}: {page_error}")
                        continue
        
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return ""
        
        return "\n\n".join(text_parts)
    
    def _ocr_pdf_page(self, page) -> str:
        """Apply OCR to a PDF page using pytesseract"""
        try:
            # Convert page to image
            image = page.to_image(resolution=300)
            
            # Use pytesseract to extract text
            text = pytesseract.image_to_string(image.original)
            return text
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DocxDocument:
            logger.error("python-docx not installed, cannot process DOCX files")
            return ""
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ""
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            return ""
    
    def _extract_image(self, file_path: str) -> str:
        """Extract text from image using OCR (pytesseract)"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""
    
    def delete_document(self, doc_id: str) -> Tuple[bool, str]:
        """
        Delete a document from RAG service
        
        Args:
            doc_id: Document ID to delete
            
        Returns:
            Tuple of (success, message)
        """
        try:
            success = self.rag_service.remove_document(doc_id)
            if success:
                return True, f"Document {doc_id} deleted"
            else:
                return False, f"Failed to delete document {doc_id}"
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False, f"Error deleting document: {str(e)}"
    
    def get_stats(self) -> Dict:
        """Get document service statistics"""
        rag_stats = self.rag_service.get_stats()
        return {
            **rag_stats,
            'supported_formats': list(self.SUPPORTED_FORMATS.keys()),
            'max_file_size_mb': self.MAX_FILE_SIZE / 1024 / 1024
        }


# Global document service instance
_document_service = None


def get_document_service() -> DocumentService:
    """Get or create document service singleton"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
